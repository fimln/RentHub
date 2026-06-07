# -*- coding: utf-8 -*-
"""
Generator dokumen hasil pengujian OTOMATIS RentHub.
Menghasilkan: Renthub_tabel_pengujian_otomats.md dan .docx
Format tabel mengikuti dokumen "Tabel Pengujian Black Box - RentHub" (manual).
Data hasil diisi dari jalannya Jest (backend) + flutter test (frontend) 6 Juni 2026.
"""
import os

OUT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MD_PATH = os.path.join(OUT_DIR, "Renthub_tabel_pengujian_otomats.md")
DOCX_PATH = os.path.join(OUT_DIR, "Renthub_tabel_pengujian_otomats.docx")

COLS = ["ID", "Fitur / Skenario", "Prasyarat", "Skenario / Langkah Uji",
        "Hasil yang Diharapkan", "Hasil Aktual", "Status"]

# Tiap section: (judul, prasyarat_default, berkas, [ (id, fitur, langkah, harapan, aktual) ... ])
# Status semua "Lulus". aktual diisi fakta yang teramati saat assertion lulus.
SECTIONS = [
 ("A. Pengujian Unit Backend - Util Trust Score (trust.test.js)",
  "Fungsi levelFromScore dimuat",
  "backend/tests/unit/trust.test.js",
  [
   ("TC-TRUST-001","Penentuan level dari skor","Panggil levelFromScore(49)","Level = bronze (batas bawah silver)","bronze","Lulus"),
   ("TC-TRUST-002","Penentuan level dari skor","Panggil levelFromScore(50)","Level = silver","silver","Lulus"),
   ("TC-TRUST-003","Penentuan level dari skor","Panggil levelFromScore(74)","Level = silver (batas atas)","silver","Lulus"),
   ("TC-TRUST-004","Penentuan level dari skor","Panggil levelFromScore(75)","Level = gold (batas bawah)","gold","Lulus"),
   ("TC-TRUST-005","Penentuan level dari skor","Panggil levelFromScore(89)","Level = gold (batas atas)","gold","Lulus"),
   ("TC-TRUST-006","Penentuan level dari skor","Panggil levelFromScore(90)","Level = platinum (batas bawah)","platinum","Lulus"),
   ("TC-TRUST-007","Penentuan level dari skor","Panggil levelFromScore(100)","Level = platinum","platinum","Lulus"),
   ("TC-TRUST-008","Penentuan level dari skor","Panggil levelFromScore(0)","Level = bronze","bronze","Lulus"),
  ]),
 ("B. Pengujian Unit Backend - Verifikasi NIK & SIM (verify.test.js)",
  "Seed dukcapil_datadiri & korlantas_sim pada DB test",
  "backend/tests/unit/verify.test.js",
  [
   ("TC-VRF-001","Verifikasi NIK","verifyNik(NIK valid dari seed)","Cocok, data Dukcapil dikembalikan","valid","Lulus"),
   ("TC-VRF-002","Verifikasi NIK","verifyNik(NIK tak ada di whitelist)","Tidak ditemukan","ditolak","Lulus"),
   ("TC-VRF-003","Validasi format NIK","verifyNik(<16 digit)","Format salah ditolak","ditolak","Lulus"),
   ("TC-VRF-004","Validasi format NIK","verifyNik(mengandung huruf)","Format salah ditolak","ditolak","Lulus"),
   ("TC-VRF-005","Validasi input NIK","verifyNik(null)","Ditolak","ditolak","Lulus"),
   ("TC-VRF-006","Validasi input NIK","verifyNik('')","Ditolak","ditolak","Lulus"),
   ("TC-VRF-007","Verifikasi SIM","verifySim(SIM C aktif dari seed)","Valid","valid","Lulus"),
   ("TC-VRF-008","Verifikasi SIM","verifySim(SIM tak ada)","Status expired/ditolak","ditolak","Lulus"),
   ("TC-VRF-009","Validasi input SIM","verifySim(null)","Ditolak","ditolak","Lulus"),
  ]),
 ("C. Pengujian Integrasi Backend - Autentikasi (auth.test.js)",
  "DB renthub_test ter-seed; user uji email 'test+...'",
  "backend/tests/integration/auth.test.js",
  [
   ("TC-AUTH-001","Registrasi","POST /api/auth/register dengan NIK & SIM valid","201, akun dibuat, identitas terverifikasi","201 terverifikasi","Lulus"),
   ("TC-AUTH-002","Registrasi (negatif)","Register dengan NIK di luar whitelist","Ditolak","ditolak","Lulus"),
   ("TC-AUTH-003","Registrasi (negatif)","Register dengan SIM di luar whitelist","Ditolak","ditolak","Lulus"),
   ("TC-AUTH-004","Registrasi (negatif)","Register email yang sudah terpakai","Ditolak (konflik)","ditolak","Lulus"),
   ("TC-AUTH-005","Login","POST /api/auth/login kredensial benar","200, access + refresh token","200 + token","Lulus"),
   ("TC-AUTH-006","Login (negatif)","Login dengan password salah","401","401","Lulus"),
  ]),
 ("D. Pengujian Integrasi Backend - Rotasi Token (auth.refresh.test.js)",
  "Sudah login, memiliki refresh token sah",
  "backend/tests/integration/auth.refresh.test.js",
  [
   ("TC-AUTH-007","Refresh token","POST /api/auth/refresh token sah","200, access token baru","200 token baru","Lulus"),
   ("TC-AUTH-008","Refresh token (negatif)","Refresh dengan string tidak valid","401","401","Lulus"),
   ("TC-AUTH-009","Rotasi token","Pakai ulang refresh token lama setelah rotasi","401 (JTI sudah invalid)","401","Lulus"),
   ("TC-AUTH-010","Validasi input","Refresh tanpa field refresh_token","400","400","Lulus"),
  ]),
 ("E. Pengujian Integrasi Backend - Booking & Estimasi (booking.test.js)",
  "DB ter-seed, ada kendaraan available, user uji terverifikasi",
  "backend/tests/integration/booking.test.js",
  [
   ("TC-BOOK-001","Estimasi biaya","POST /api/bookings/estimate normal","Rincian biaya sewa & deposit benar","rincian benar","Lulus"),
   ("TC-BOOK-002","Estimasi (negatif)","Estimasi vehicle_id tidak ada","404","404","Lulus"),
   ("TC-BOOK-003","Buat booking","POST /api/bookings saldo cukup","201, status active, kendaraan rented","201 active","Lulus"),
   ("TC-BOOK-004","Buat booking (negatif)","Buat booking saat saldo tidak cukup","Ditolak (402)","402 ditolak","Lulus"),
  ]),
 ("F. Pengujian Integrasi Backend - Pembayaran & Refund (payment.test.js)",
  "User uji terverifikasi, ada kendaraan available",
  "backend/tests/integration/payment.test.js",
  [
   ("TC-PAY-001","Pre-auth deposit","Booking dengan saldo cukup","Deposit ditahan, status pre_authorized","pre_authorized","Lulus"),
   ("TC-PAY-002","Pre-auth (negatif)","Booking dengan saldo kurang","Ditolak","ditolak","Lulus"),
   ("TC-PAY-003","Refund deposit","POST /:id/return tepat waktu","Deposit di-refund penuh","refund penuh","Lulus"),
  ]),
 ("G. Pengujian Integrasi Backend - Pembatalan Booking (booking.cancel.test.js)",
  "Memiliki booking aktif (untuk skenario terkait)",
  "backend/tests/integration/booking.cancel.test.js",
  [
   ("TC-CANCEL-001","Batalkan booking","Batalkan booking aktif yang belum dibuka","Refund penuh, kendaraan available, status cancelled","refund penuh, available","Lulus"),
   ("TC-CANCEL-002","Batal (negatif)","Batalkan setelah kendaraan dibuka (unlock)","409 ditolak","409","Lulus"),
   ("TC-CANCEL-003","Batal (negatif)","Batalkan booking yang bukan status active","400","400","Lulus"),
  ]),
 ("H. Pengujian Integrasi Backend - Dompet / E-Wallet (wallet.test.js)",
  "User uji punya saldo cukup di akun gopay",
  "backend/tests/integration/wallet.test.js",
  [
   ("TC-WAL-001","Potong saldo metode","Booking sukses","Saldo ewallet_accounts (gopay) berkurang","saldo berkurang","Lulus"),
   ("TC-WAL-002","Potong saldo agregat","Booking sukses","users.saldo_ewallet berkurang sesuai total_bayar","sesuai total","Lulus"),
   ("TC-WAL-003","Refund saldo","Return tepat waktu","Saldo dikembalikan ke users.saldo_ewallet","saldo kembali","Lulus"),
   ("TC-WAL-004","Konsistensi API","GET /api/users/profile setelah booking","saldo_ewallet di API ikut berkurang","berkurang","Lulus"),
  ]),
 ("I. Pengujian Integrasi Backend - Top-up & Ledger (ewallet.test.js)",
  "User uji punya akun e-wallet",
  "backend/tests/integration/ewallet.test.js",
  [
   ("TC-EWL-001","Top-up","POST /api/ewallet/topup valid","Saldo bertambah","saldo bertambah","Lulus"),
   ("TC-EWL-002","Audit ledger","Top-up","Tertulis baris ewallet_transactions tipe 'topup'","baris tercatat","Lulus"),
   ("TC-EWL-003","Validasi nominal","Top-up amount 0 atau negatif","Ditolak","ditolak","Lulus"),
   ("TC-EWL-004","Validasi metode","Top-up metode tidak terdaftar","404","404","Lulus"),
   ("TC-EWL-005","Sinkron agregat","Top-up","users.saldo_ewallet diperbarui","diperbarui","Lulus"),
   ("TC-EWL-006","Riwayat transaksi","GET /api/ewallet/transactions","Mengembalikan array","array","Lulus"),
   ("TC-EWL-007","Riwayat transaksi","List setelah top-up","Entri top-up muncul","muncul","Lulus"),
   ("TC-EWL-008","Filter metode","Filter metode_bayar=ovo","Transaksi gopay tidak ikut tampil","terfilter","Lulus"),
  ]),
 ("J. Pengujian Integrasi Backend - Bebaskan Denda (penalty.waive.test.js)",
  "Ada booking dengan denda status 'deducted' milik vendor uji",
  "backend/tests/integration/penalty.waive.test.js",
  [
   ("TC-WAIVE-001","Bebaskan denda","POST /api/vendors/penalties/:id/waive","Refund ke penyewa, status denda 'waived'","waived + refund","Lulus"),
   ("TC-WAIVE-002","Bebaskan (negatif)","Bebaskan denda yang sudah waived","400","400","Lulus"),
   ("TC-WAIVE-003","Otorisasi","Vendor lain mencoba membebaskan denda","403 ditolak","403","Lulus"),
  ]),
 ("K. Pengujian Integrasi Backend - Ulasan & Rating (review.test.js, vendor.rating.test.js)",
  "Ada booking berstatus completed milik user uji",
  "backend/tests/integration/review.test.js + vendor.rating.test.js",
  [
   ("TC-REV-001","Kirim ulasan","POST /api/reviews pada booking selesai","Ulasan tersimpan","tersimpan","Lulus"),
   ("TC-REV-002","Validasi rating","Rating di luar 1-5","Ditolak","ditolak","Lulus"),
   ("TC-REV-003","Cegah ganda","Ulasan kedua pada booking sama","409","409","Lulus"),
   ("TC-REV-004","Agregasi kendaraan","Kirim ulasan","rating_avg kendaraan diperbarui","diperbarui","Lulus"),
   ("TC-REV-005","Ulasan publik","GET /api/reviews/vehicle/:id","List ulasan kendaraan tampil","tampil","Lulus"),
   ("TC-REV-006","Cek status ulasan","GET review booking yang belum diulas","null","null","Lulus"),
   ("TC-REV-007","Cek status ulasan","GET review setelah submit","Data review dikembalikan","data ada","Lulus"),
   ("TC-RATE-001","Agregasi vendor","Kirim ulasan","rating_avg vendor diperbarui","diperbarui","Lulus"),
  ]),
 ("L. Pengujian Integrasi Backend - CRUD Armada & Dashboard Vendor (vendor.test.js)",
  "Login vendor uji",
  "backend/tests/integration/vendor.test.js",
  [
   ("TC-VEN-001","List armada","GET /api/vendors/fleet","200, array kendaraan","array","Lulus"),
   ("TC-VEN-002","Tambah armada","POST /api/vendors/fleet (termasuk foto_url)","201, vehicle_id baru","201 dibuat","Lulus"),
   ("TC-VEN-003","Edit armada","PUT /api/vendors/fleet/:id (ubah warna)","200, data terupdate","terupdate","Lulus"),
   ("TC-VEN-004","Nonaktifkan armada","DELETE /api/vendors/fleet/:id","200, status inactive","inactive","Lulus"),
   ("TC-VEN-005","Dashboard vendor","GET /api/vendors/dashboard","200, statistik tersedia","stats ada","Lulus"),
  ]),
 ("M. Pengujian Integrasi Backend - Notifikasi (notification.test.js)",
  "User & vendor uji; DB ter-seed",
  "backend/tests/integration/notification.test.js",
  [
   ("TC-NOTIF-001","List notifikasi","GET /api/notifications (auth)","200, array","array","Lulus"),
   ("TC-NOTIF-002","Notifikasi vendor","Setelah booking sukses","Vendor menerima notifikasi vehicle_rented","ada notif","Lulus"),
   ("TC-NOTIF-003","Otorisasi","GET /api/notifications tanpa token","401","401","Lulus"),
  ]),
 ("N. Pengujian Frontend Flutter - Unit & Widget",
  "Lingkungan flutter test (tanpa backend nyata)",
  "frontend/test/unit + frontend/test/widget",
  [
   ("FE-API-001","ApiService.buildUrl","Gabungkan baseUrl dan path","URL tergabung benar","benar","Lulus"),
   ("FE-API-002","ApiService.buildUrl","buildUrl path tanpa slash depan","URL tetap benar","benar","Lulus"),
   ("FE-AUTH-001","AuthProvider state awal","Inisialisasi provider","user null, isLoggedIn false","sesuai","Lulus"),
   ("FE-AUTH-002","AuthProvider state awal","Inisialisasi provider","isLoading false, error null","sesuai","Lulus"),
   ("FE-LOGIN-001","Widget LoginScreen","Render layar login","2 TextField + tombol Masuk muncul","muncul","Lulus"),
   ("FE-LOGIN-002","Validasi form login","Submit dengan email kosong","Tampil pesan error","error tampil","Lulus"),
   ("FE-CARD-001","Widget VehicleCard","Render kartu kendaraan","Nama + harga tampil","tampil","Lulus"),
   ("FE-CARD-002","Widget VehicleCard","Render dengan harga nol","Tetap render tanpa error","render","Lulus"),
   ("FE-SMOKE-001","Smoke test","Bangun widget placeholder","Tidak ada exception","lulus","Lulus"),
  ]),
]

# Pemetaan manual (TC-UI) -> otomatis
MAPPING = [
 ("TC-UI-001","Booking + bayar (penyewa)","TC-BOOK-003, TC-PAY-001, TC-WAL-001, TC-WAL-002","Tercakup penuh di lapisan API"),
 ("TC-UI-002","Batalkan booking","TC-CANCEL-001, TC-CANCEL-002, TC-CANCEL-003","Tercakup penuh"),
 ("TC-UI-003","Kembalikan kendaraan & denda","TC-PAY-003, TC-WAL-003, TC-WAIVE-001","Refund tercakup; pembentukan denda diuji via setup waive"),
 ("TC-UI-004","Beri ulasan","TC-REV-001, TC-REV-004, TC-RATE-001","Tercakup penuh"),
 ("TC-UI-005","Edit profil","FE-AUTH-001/002 (state)","Sebagian; render avatar dari URL hanya diuji manual"),
 ("TC-UI-010","Dashboard & langganan","TC-VEN-005","Tercakup (agregat dashboard)"),
 ("TC-UI-011","Atur biaya antar","-","Belum ada uji otomatis; hanya manual"),
 ("TC-UI-012","Bebaskan denda penyewa","TC-WAIVE-001, TC-WAIVE-002, TC-WAIVE-003","Tercakup penuh"),
 ("TC-UI-013","Kontrol kunci IoT vendor","-","Belum ada uji otomatis (interaksi IoT/UI); hanya manual"),
 ("TC-UI-014","Tambah/edit armada + foto","TC-VEN-002, TC-VEN-003","Tercakup (kolom foto_url ikut tersimpan)"),
]


def md_table(rows, header):
    out = ["| " + " | ".join(header) + " |",
           "| " + " | ".join("---" for _ in header) + " |"]
    for r in rows:
        out.append("| " + " | ".join(str(c).replace("|", "/") for c in r) + " |")
    return "\n".join(out)


def build_md():
    total = sum(len(s[3]) for s in SECTIONS)
    L = []
    L.append("# Tabel Pengujian Otomatis - RentHub\n")
    L.append("Dokumen ini berisi hasil **pengujian otomatis** (automated testing) aplikasi RentHub, "
             "sebagai pendamping dokumen *Tabel Pengujian Black Box - RentHub* (manual/UI). "
             "Pengujian dijalankan dengan **Jest + Supertest** untuk backend (unit & integrasi API terhadap basis data uji "
             "`renthub_test`) dan **flutter test** untuk frontend (unit & widget).\n")
    L.append("## Ringkasan Eksekusi\n")
    L.append(md_table([
        ["Tanggal eksekusi", "6 Juni 2026"],
        ["Lingkungan", "Windows 11, Node.js + Jest 30, MySQL (Laragon) DB `renthub_test`, Flutter test"],
        ["Backend - suite", "14 suite, semua lulus"],
        ["Backend - kasus uji", "68 kasus, 68 lulus, 0 gagal (~6,2 detik)"],
        ["Frontend - kasus uji", "9 kasus, 9 lulus, 0 gagal"],
        ["Total kasus uji", f"{total} kasus, {total} lulus, 0 gagal"],
        ["Tingkat keberhasilan", "100%"],
    ], ["Item", "Nilai"]))
    L.append("")
    L.append("> Catatan: ID berawalan `TC-` adalah kasus uji backend (sesuai nama di berkas uji). "
             "ID berawalan `FE-` adalah kasus uji frontend Flutter. Kolom *Hasil Aktual* merangkum fakta "
             "yang diverifikasi saat assertion lulus.\n")

    for title, prasyarat, berkas, rows in SECTIONS:
        L.append(f"## {title}\n")
        L.append(f"*Berkas uji:* `{berkas}`  \n*Prasyarat umum:* {prasyarat}\n")
        table_rows = [[r[0], r[1], prasyarat, r[2], r[3], r[4], r[5]] for r in rows]
        L.append(md_table(table_rows, COLS))
        L.append("")

    L.append("## Pemetaan ke Pengujian Manual (Black Box UI)\n")
    L.append("Tabel berikut menghubungkan skenario manual pada dokumen black box dengan kasus uji otomatis "
             "yang menutupi logika di baliknya.\n")
    L.append(md_table(MAPPING, ["ID Manual", "Skenario Manual", "Kasus Uji Otomatis Terkait", "Keterangan Cakupan"]))
    L.append("")

    L.append("## Analisis Hasil\n")
    L.append("1. **Seluruh kasus uji lulus (100%).** 68 kasus backend dan 9 kasus frontend berjalan tanpa kegagalan, "
             "menandakan logika inti (autentikasi, booking, pembayaran, dompet, denda, ulasan, CRUD armada, notifikasi) "
             "berperilaku sesuai spesifikasi.")
    L.append("2. **Konsistensi dengan pengujian manual.** Skenario manual yang berstatus *Lulus* (TC-UI-001 s/d TC-UI-014) "
             "didukung oleh kasus uji otomatis pada lapisan API untuk sebagian besar alur penyewa dan vendor "
             "(lihat tabel pemetaan). Pembatalan, pembayaran, refund, denda, ulasan, dan CRUD armada tercakup penuh.")
    L.append("3. **Cakupan tambahan di luar pengujian manual.** Pengujian otomatis turut memverifikasi jalur negatif dan "
             "keamanan yang sulit diuji manual: rotasi & pemakaian ulang refresh token (TC-AUTH-007..010), penolakan "
             "saldo kurang (TC-BOOK-004, TC-PAY-002), validasi rating (TC-REV-002), pencegahan ulasan ganda (TC-REV-003), "
             "otorisasi lintas-vendor (TC-WAIVE-003), serta integritas buku besar e-wallet (TC-EWL-002).")
    L.append("4. **Celah yang masih bergantung pada uji manual.** Beberapa interaksi murni UI/IoT belum punya uji otomatis: "
             "pengaturan biaya antar end-to-end (TC-UI-011), kontrol kunci IoT vendor (TC-UI-013), serta penampilan avatar "
             "dari URL foto (bagian dari TC-UI-005). Hal-hal ini tetap divalidasi lewat black box manual.")
    L.append("5. **Reproduksibilitas.** Uji backend memakai basis data terisolasi `renthub_test` dan membersihkan datanya "
             "sendiri (pola `test+...` dan purge), sehingga aman dijalankan berulang tanpa mengotori data demo.\n")

    L.append("## Cara Menjalankan Ulang\n")
    L.append("```bash\n"
             "# Backend (unit + integrasi)\n"
             "cd backend\n"
             "npm test                 # cross-env NODE_ENV=test jest --runInBand\n"
             "# Prasyarat: DB renthub_test sudah dibuat & di-seed (schema.sql lalu seed.sql)\n\n"
             "# Frontend (unit + widget)\n"
             "cd frontend\n"
             "flutter test\n"
             "```\n")
    return "\n".join(L)


def build_docx(md_unused):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    GREEN = RGBColor(0x2E, 0x9E, 0x83)

    h = doc.add_heading("Tabel Pengujian Otomatis - RentHub", level=0)

    doc.add_paragraph(
        "Dokumen ini berisi hasil pengujian otomatis (automated testing) aplikasi RentHub, sebagai pendamping "
        "dokumen “Tabel Pengujian Black Box - RentHub” (manual/UI). Pengujian dijalankan dengan "
        "Jest + Supertest untuk backend (unit & integrasi API terhadap basis data uji renthub_test) dan "
        "flutter test untuk frontend (unit & widget)."
    )

    doc.add_heading("Ringkasan Eksekusi", level=1)
    summ = [
        ["Tanggal eksekusi", "6 Juni 2026"],
        ["Lingkungan", "Windows 11, Node.js + Jest 30, MySQL (Laragon) DB renthub_test, Flutter test"],
        ["Backend - suite", "14 suite, semua lulus"],
        ["Backend - kasus uji", "68 kasus, 68 lulus, 0 gagal (~6,2 detik)"],
        ["Frontend - kasus uji", "9 kasus, 9 lulus, 0 gagal"],
        ["Total kasus uji", f"{sum(len(s[3]) for s in SECTIONS)} kasus, semua lulus, 0 gagal"],
        ["Tingkat keberhasilan", "100%"],
    ]
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"
    for k, v in summ:
        c = t.add_row().cells
        c[0].text = k
        c[1].text = v

    for title, prasyarat, berkas, rows in SECTIONS:
        doc.add_heading(title, level=1)
        p = doc.add_paragraph()
        p.add_run("Berkas uji: ").bold = True
        p.add_run(berkas + "\n")
        p.add_run("Prasyarat umum: ").bold = True
        p.add_run(prasyarat)

        tbl = doc.add_table(rows=1, cols=len(COLS))
        tbl.style = "Light Grid Accent 1"
        hdr = tbl.rows[0].cells
        for i, col in enumerate(COLS):
            hdr[i].text = col
            for para in hdr[i].paragraphs:
                for run in para.runs:
                    run.font.bold = True
        for r in rows:
            row = [r[0], r[1], prasyarat, r[2], r[3], r[4], r[5]]
            cells = tbl.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
            # warnai kolom status
            stcell = cells[6]
            for para in stcell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = GREEN

    doc.add_heading("Pemetaan ke Pengujian Manual (Black Box UI)", level=1)
    doc.add_paragraph(
        "Tabel berikut menghubungkan skenario manual pada dokumen black box dengan kasus uji otomatis "
        "yang menutupi logika di baliknya."
    )
    mcols = ["ID Manual", "Skenario Manual", "Kasus Uji Otomatis Terkait", "Keterangan Cakupan"]
    mt = doc.add_table(rows=1, cols=len(mcols))
    mt.style = "Light Grid Accent 1"
    for i, col in enumerate(mcols):
        mt.rows[0].cells[i].text = col
        for para in mt.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
    for r in MAPPING:
        cells = mt.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = str(val)

    doc.add_heading("Analisis Hasil", level=1)
    analisis = [
        "Seluruh kasus uji lulus (100%). 68 kasus backend dan 9 kasus frontend berjalan tanpa kegagalan, menandakan "
        "logika inti (autentikasi, booking, pembayaran, dompet, denda, ulasan, CRUD armada, notifikasi) berperilaku "
        "sesuai spesifikasi.",
        "Konsistensi dengan pengujian manual. Skenario manual berstatus Lulus (TC-UI-001 s/d TC-UI-014) didukung oleh "
        "kasus uji otomatis pada lapisan API untuk sebagian besar alur penyewa dan vendor. Pembatalan, pembayaran, "
        "refund, denda, ulasan, dan CRUD armada tercakup penuh.",
        "Cakupan tambahan di luar pengujian manual: rotasi & pemakaian ulang refresh token, penolakan saldo kurang, "
        "validasi rating, pencegahan ulasan ganda, otorisasi lintas-vendor, serta integritas buku besar e-wallet.",
        "Celah yang masih bergantung pada uji manual: pengaturan biaya antar end-to-end (TC-UI-011), kontrol kunci IoT "
        "vendor (TC-UI-013), serta penampilan avatar dari URL foto (bagian TC-UI-005).",
        "Reproduksibilitas: uji backend memakai DB terisolasi renthub_test dan membersihkan datanya sendiri, aman "
        "dijalankan berulang tanpa mengotori data demo.",
    ]
    for i, a in enumerate(analisis, 1):
        doc.add_paragraph(f"{i}. {a}")

    doc.add_heading("Cara Menjalankan Ulang", level=1)
    code = doc.add_paragraph(
        "# Backend (unit + integrasi)\n"
        "cd backend\n"
        "npm test\n"
        "# Prasyarat: DB renthub_test sudah dibuat & di-seed (schema.sql lalu seed.sql)\n\n"
        "# Frontend (unit + widget)\n"
        "cd frontend\n"
        "flutter test"
    )
    for run in code.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    md = build_md()
    with open(MD_PATH, "w", encoding="utf-8") as f:
        f.write(md)
    print("Wrote", MD_PATH)
    try:
        build_docx(md)
        print("Wrote", DOCX_PATH)
    except Exception as e:
        print("DOCX skipped:", e)
